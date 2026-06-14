# -*- coding: utf-8 -*-
"""Generate Vietnamese user guide DOCX for KOL Hub."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT = os.path.join(
    os.path.dirname(os.path.dirname(__file__)),
    "docs",
    "Huong-dan-su-dung-KOL-Hub.docx",
)


def set_cell_shading(cell, color_hex: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    return h


def add_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run("💡 Lưu ý: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0xB4, 0x53, 0x09)
    p.add_run(text)
    p.paragraph_format.left_indent = Cm(0.5)
    return p


def add_steps(doc, steps):
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(step)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], "1A56DB")
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = val
    doc.add_paragraph()
    return table


def build_doc():
    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    # Cover
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("HƯỚNG DẪN SỬ DỤNG\nKOL HUB")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Nền tảng kết nối Thương hiệu và Nhà sáng tạo (KOL)")
    r.font.size = Pt(14)
    r.italic = True

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run("Phiên bản: 1.0  |  Ngôn ngữ: Tiếng Việt\n")
    info.add_run("Dành cho: Quản trị viên, Thương hiệu, Nhà sáng tạo (KOL)")

    doc.add_page_break()

    # Mục lục nhanh
    add_heading(doc, "Mục lục", 1)
    toc_items = [
        "1. Giới thiệu chung",
        "2. Bắt đầu: Đăng ký và đăng nhập",
        "3. Hướng dẫn cho Thương hiệu (Brand)",
        "4. Hướng dẫn cho Nhà sáng tạo (KOL)",
        "5. Hướng dẫn cho Quản trị viên (Admin)",
        "6. Các tính năng dùng chung",
        "7. Câu hỏi thường gặp (FAQ)",
        "8. Bảng thuật ngữ",
    ]
    add_bullets(doc, toc_items)
    doc.add_page_break()

    # 1. Giới thiệu
    add_heading(doc, "1. Giới thiệu chung", 1)
    doc.add_paragraph(
        "KOL Hub là website giúp Thương hiệu (công ty, doanh nghiệp) tìm và "
        "hợp tác với các Nhà sáng tạo nội dung (KOL – Key Opinion Leader) "
        "trên mạng xã hội như TikTok, Instagram, YouTube, Facebook."
    )
    doc.add_paragraph(
        "Trên website có 3 loại tài khoản chính:"
    )
    add_table(doc, ["Vai trò", "Ai sử dụng?", "Làm gì trên website?"], [
        ["Thương hiệu (Brand)", "Công ty, doanh nghiệp, marketer", "Tìm KOL, đăng chiến dịch, đặt KOL, thanh toán"],
        ["Nhà sáng tạo (KOL)", "Người làm nội dung, influencer", "Tạo hồ sơ, ứng tuyển chiến dịch, nhận đơn, rút tiền"],
        ["Quản trị viên (Admin)", "Nhân viên vận hành hệ thống", "Duyệt hồ sơ, quản lý người dùng, xử lý rút tiền"],
    ])
    add_note(doc, "Tài khoản Admin không tự đăng ký được. Admin được cấp tài khoản bởi đội kỹ thuật.")

    # 2. Đăng ký đăng nhập
    add_heading(doc, "2. Bắt đầu: Đăng ký và đăng nhập", 1)

    add_heading(doc, "2.1. Đăng ký tài khoản mới", 2)
    add_steps(doc, [
        "Mở website KOL Hub, bấm nút \"Đăng ký\" ở góc trên bên phải.",
        "Chọn loại tài khoản:\n   • Thương hiệu — nếu bạn là công ty muốn tìm KOL\n   • Nhà sáng tạo — nếu bạn là KOL muốn nhận việc",
        "Nhập email, mật khẩu (tối thiểu 8 ký tự) và xác nhận mật khẩu.",
        "Bấm \"Tiếp tục\".",
        "Kiểm tra email — bạn sẽ nhận được thư xác nhận. Mở email và bấm vào liên kết xác nhận.",
        "Sau khi xác nhận thành công, website sẽ tự chuyển bạn vào trang chính của tài khoản.",
    ])
    add_note(doc, "Không thấy email? Kiểm tra thư mục Spam/Quảng cáo. Bạn có thể bấm \"Gửi lại email xác nhận\" (chờ 60 giây giữa mỗi lần gửi).")

    add_heading(doc, "2.2. Đăng nhập", 2)
    add_steps(doc, [
        "Bấm \"Đăng nhập\" ở góc trên bên phải.",
        "Nhập email và mật khẩu đã đăng ký.",
        "Bấm đăng nhập — website tự đưa bạn đến trang phù hợp với vai trò của bạn.",
    ])

    add_heading(doc, "2.3. Đăng xuất", 2)
    doc.add_paragraph(
        "Bấm vào ảnh đại diện (avatar) ở góc trên bên phải, chọn \"Đăng xuất\"."
    )

    add_heading(doc, "2.4. Thanh menu chung", 2)
    doc.add_paragraph("Dù bạn là ai, khi vào website bạn sẽ thấy các mục sau trên thanh menu:")
    add_table(doc, ["Tên trên menu", "Dùng để làm gì?"], [
        ["Khám phá", "Tìm kiếm và lọc danh sách KOL"],
        ["Hồ sơ KOL", "Xem danh sách hồ sơ KOL công khai"],
        ["Chiến dịch", "Xem các chiến dịch đang tuyển KOL"],
        ["AI KOL", "Trò chuyện với trợ lý AI để được gợi ý KOL phù hợp"],
        ["Đơn đặt", "Xem danh sách đơn hàng (Brand và KOL)"],
        ["Quản trị", "Chỉ Admin mới thấy — vào khu vực quản trị"],
    ])
    doc.add_page_break()

    # 3. BRAND
    add_heading(doc, "3. Hướng dẫn cho Thương hiệu (Brand)", 1)
    doc.add_paragraph(
        "Phần này dành cho bạn nếu bạn đại diện cho một công ty hoặc thương hiệu "
        "muốn tìm KOL để quảng bá sản phẩm/dịch vụ."
    )

    add_heading(doc, "3.1. Trang Bảng điều khiển", 2)
    doc.add_paragraph(
        "Sau khi đăng nhập, bạn vào trang \"Bảng điều khiển\". Tại đây bạn thấy:"
    )
    add_bullets(doc, [
        "Tổng số đơn đặt của bạn",
        "Số dư ví",
        "Tổng ngân sách đã chi",
        "Điểm đánh giá trung bình",
        "Các nút tắt nhanh: Đăng sản phẩm, Tin đăng của tôi, Khám phá KOL, Đơn đặt",
    ])

    add_heading(doc, "3.2. Hoàn thiện hồ sơ công ty (BẮT BUỘC trước khi đặt KOL)", 2)
    doc.add_paragraph(
        "Trước khi đặt KOL hoặc đăng chiến dịch, bạn cần điền đầy đủ thông tin công ty "
        "và chờ Admin duyệt."
    )
    add_steps(doc, [
        "Bấm vào ảnh đại diện → chọn \"Hồ sơ\" (hoặc vào menu \"Hồ sơ của tôi\").",
        "Tải lên ảnh đại diện (logo công ty).",
        "Điền thông tin cá nhân: Họ tên, Số điện thoại, Quốc gia, Địa chỉ, Giới thiệu ngắn.",
        "Điền thông tin công ty: Tên công ty, Ngành nghề.",
        "Bấm \"Lưu thay đổi\".",
        "Chờ Admin duyệt hồ sơ. Trạng thái sẽ hiển thị \"Chờ duyệt\" → \"Đã duyệt\".",
    ])
    add_note(doc, "Bạn chỉ có thể đặt KOL khi hồ sơ đã được duyệt (trạng thái \"Đã duyệt\").")

    add_heading(doc, "3.3. Đăng chiến dịch tuyển KOL", 2)
    doc.add_paragraph(
        "Nếu bạn muốn KOL tự ứng tuyển vào chiến dịch của mình (thay vì bạn chủ động chọn KOL):"
    )
    add_steps(doc, [
        "Vào Bảng điều khiển → bấm \"Đăng sản phẩm\" (hoặc menu \"Đăng sản phẩm\").",
        "Điền thông tin chiến dịch: tên, mô tả, ngân sách, thời gian, yêu cầu...",
        "Bấm \"Đăng tin\" để đăng chiến dịch lên website.",
        "Quản lý tin đăng tại \"Tin đăng của tôi\": có thể sửa, mở tin hoặc khóa tin.",
    ])

    add_heading(doc, "3.4. Xem và duyệt ứng viên KOL", 2)
    add_steps(doc, [
        "Vào \"Tin đăng của tôi\" → chọn chiến dịch → bấm \"Ứng viên\".",
        "Xem danh sách KOL đã ứng tuyển. Bạn có thể lọc theo trạng thái: Chờ duyệt, Shortlist, Đã duyệt, Từ chối.",
        "Với ứng viên phù hợp: bấm \"Shortlist\" để đánh dấu ưu tiên.",
        "Khi quyết định chọn KOL: bấm \"Duyệt\" — hệ thống tự tạo đơn đặt.",
        "Nếu không chọn: bấm \"Từ chối\" và có thể ghi lý do.",
    ])

    add_heading(doc, "3.5. Tìm và đặt KOL trực tiếp", 2)
    add_steps(doc, [
        "Vào \"Khám phá\" hoặc \"Hồ sơ KOL\" để tìm KOL.",
        "Dùng bộ lọc (danh mục, nền tảng, ngân sách...) hoặc ô tìm kiếm.",
        "Bấm vào hồ sơ KOL để xem chi tiết: kênh mạng xã hội, gói giá, portfolio.",
        "Bấm \"Đặt KOL này\".",
        "Điền form: tên chiến dịch, mô tả (tối thiểu 50 ký tự), yêu cầu sản phẩm, ngân sách, ngày bắt đầu và kết thúc.",
        "Bấm \"Gửi yêu cầu\" — đơn sẽ ở trạng thái \"Chờ duyệt\" (chờ KOL chấp nhận).",
    ])

    add_heading(doc, "3.6. Quy trình đơn đặt (từ lúc gửi đến khi hoàn thành)", 2)
    doc.add_paragraph("Đơn đặt sẽ trải qua các bước sau:")
    add_table(doc, ["Trạng thái", "Ý nghĩa", "Bạn cần làm gì?"], [
        ["Chờ duyệt", "Đã gửi yêu cầu, chờ KOL phản hồi", "Chờ KOL chấp nhận, hoặc bấm \"Hủy đơn\" nếu muốn hủy"],
        ["Đã chấp nhận", "KOL đồng ý nhận việc", "Bấm \"Thanh toán\" — chuyển sang cổng VNPay để thanh toán"],
        ["Đang thực hiện", "Đã thanh toán, KOL đang làm nội dung", "Theo dõi tiến độ, trò chuyện với KOL trong tab \"Trò chuyện\""],
        ["Đã giao", "KOL đã nộp sản phẩm (link bài đăng/video)", "Kiểm tra nội dung → bấm \"Xác nhận hoàn thành\" nếu OK"],
        ["Hoàn thành", "Chiến dịch kết thúc, tiền đã chuyển cho KOL", "Có thể để lại đánh giá, tải hóa đơn"],
        ["Tranh chấp", "Có vấn đề cần xử lý", "Bấm \"Báo cáo vấn đề\" để Admin hỗ trợ"],
    ])
    add_note(doc, "Tiền thanh toán được giữ an toàn trên hệ thống cho đến khi bạn xác nhận KOL đã hoàn thành công việc.")

    add_heading(doc, "3.7. Thanh toán qua VNPay", 2)
    add_steps(doc, [
        "Khi đơn ở trạng thái \"Đã chấp nhận\", vào chi tiết đơn → bấm \"Thanh toán\".",
        "Website chuyển bạn sang trang VNPay.",
        "Chọn phương thức thanh toán (thẻ, ví điện tử, chuyển khoản...) và hoàn tất.",
        "Sau thanh toán, bạn được chuyển về website với thông báo thành công hoặc thất bại.",
    ])

    add_heading(doc, "3.8. Dùng trợ lý AI", 2)
    doc.add_paragraph(
        "Vào menu \"AI KOL\" để trò chuyện với trợ lý AI. Bạn mô tả chiến dịch "
        "(sản phẩm, ngân sách, đối tượng khách hàng...) và AI sẽ gợi ý KOL phù hợp. "
        "Cần đăng nhập mới chat được."
    )
    doc.add_page_break()

    # 4. KOL
    add_heading(doc, "4. Hướng dẫn cho Nhà sáng tạo (KOL)", 1)
    doc.add_paragraph(
        "Phần này dành cho bạn nếu bạn là người làm nội dung trên mạng xã hội "
        "và muốn nhận các cơ hội hợp tác từ thương hiệu."
    )

    add_heading(doc, "4.1. Trang quản lý KOL", 2)
    doc.add_paragraph(
        "Sau đăng nhập, bạn vào \"Trang quản lý KOL\" với các tab:"
    )
    add_bullets(doc, [
        "Tổng quan — thống kê nhanh, đơn chờ duyệt",
        "Đơn đặt — danh sách đơn hàng",
        "Nhận xét — đánh giá từ thương hiệu",
        "Cài đặt — tùy chỉnh tài khoản",
    ])

    add_heading(doc, "4.2. Tạo và gửi hồ sơ KOL (BẮT BUỘC)", 2)
    doc.add_paragraph(
        "Hồ sơ KOL giống như \"CV online\" của bạn trên nền tảng. "
        "Bạn cần hoàn thiện và gửi duyệt trước khi ứng tuyển hoặc nhận đơn."
    )
    add_steps(doc, [
        "Vào \"Trang quản lý KOL\" → bấm \"Chỉnh sửa hồ sơ\".",
        "Điền thông tin cơ bản: Tên hiển thị, mô tả bản thân (bio), giới tính, ngày sinh, thành phố.",
        "Tải ảnh đại diện và ảnh bìa.",
        "Chọn danh mục phù hợp (làm đẹp, thời trang, ẩm thực...).",
        "Thêm kênh mạng xã hội: chọn nền tảng (TikTok, Instagram, YouTube, Facebook), nhập link và số follower.",
        "Thêm gói giá: loại nội dung (Bài đăng, Story, Video...) và mức giá.",
        "Thêm portfolio: ảnh hoặc link các bài đăng tiêu biểu.",
        "Khi đã đủ (ít nhất 1 kênh + 1 gói giá + 1 portfolio), bấm \"Gửi duyệt\".",
        "Chờ Admin duyệt. Trạng thái: Bản nháp → Chờ duyệt → Đã duyệt (hoặc Bị từ chối kèm lý do).",
    ])
    add_note(doc, "Chỉ khi hồ sơ \"Đã duyệt\" bạn mới ứng tuyển chiến dịch và nhận đơn đặt trực tiếp.")

    add_heading(doc, "4.3. Ứng tuyển chiến dịch", 2)
    add_steps(doc, [
        "Vào menu \"Chiến dịch\" để xem danh sách chiến dịch đang tuyển KOL.",
        "Dùng bộ lọc theo danh mục, nền tảng, ngân sách.",
        "Bấm vào chiến dịch quan tâm → đọc kỹ yêu cầu.",
        "Gửi ứng tuyển: viết lời giới thiệu ngắn, đề xuất giá (nếu có).",
        "Theo dõi tại \"Ứng tuyển của tôi\". Có thể \"Rút ứng tuyển\" nếu đổi ý.",
        "Nếu Brand duyệt bạn → hệ thống tự tạo đơn đặt.",
    ])

    add_heading(doc, "4.4. Nhận và xử lý đơn đặt trực tiếp", 2)
    doc.add_paragraph("Khi Brand đặt bạn trực tiếp, xem đơn tại \"Đơn đặt\" (Đơn được gửi đến):")
    add_table(doc, ["Trạng thái", "Bạn cần làm gì?"], [
        ["Chờ duyệt", "Xem chi tiết → \"Chấp nhận\" hoặc \"Từ chối\" (có thể ghi lý do)"],
        ["Đang thực hiện", "Brand đã thanh toán — làm nội dung theo yêu cầu"],
        ["Đang thực hiện", "Khi xong: bấm \"Nộp deliverable\" — dán link bài đăng/video và ghi chú"],
        ["Hoàn thành", "Brand đã xác nhận — tiền vào ví. Có thể để lại đánh giá cho Brand"],
    ])

    add_heading(doc, "4.5. Ví KOL và rút tiền", 2)
    add_steps(doc, [
        "Vào \"Trang quản lý KOL\" → \"Ví KOL\".",
        "Xem số dư: \"Khả dụng\" (rút được) và \"Tạm giữ\" (đang chờ Brand xác nhận).",
        "Xem lịch sử giao dịch: Nạp tiền, Tạm giữ, Giải ngân, Rút tiền, Hoàn tiền, Phí nền tảng.",
        "Để rút tiền: nhập số tiền + thông tin ngân hàng (tên ngân hàng, số tài khoản, tên chủ tài khoản).",
        "Gửi yêu cầu rút tiền → chờ Admin duyệt và chuyển khoản.",
        "Theo dõi trạng thái: Chờ duyệt → Đã duyệt → Đã chi trả (hoặc Đã từ chối).",
    ])
    add_note(doc, "Hệ thống thu một khoản phí nền tảng (hoa hồng) trên mỗi giao dịch thành công. Số tiền thực nhận = ngân sách đơn hàng trừ phí.")

    add_heading(doc, "4.6. Xem hồ sơ công khai", 2)
    doc.add_paragraph(
        "Bấm \"Xem hồ sơ công khai\" hoặc \"Xem trước hồ sơ\" để xem hồ sơ "
        "như Brand sẽ thấy khi tìm KOL."
    )
    doc.add_page_break()

    # 5. ADMIN
    add_heading(doc, "5. Hướng dẫn cho Quản trị viên (Admin)", 1)
    doc.add_paragraph(
        "Phần này dành cho nhân viên vận hành nền tảng. "
        "Admin vào khu vực riêng qua menu \"Quản trị\" hoặc đường dẫn /admin."
    )

    add_heading(doc, "5.1. Menu quản trị", 2)
    add_table(doc, ["Menu", "Chức năng"], [
        ["Tổng quan", "Xem thống kê toàn nền tảng: người dùng, doanh thu, booking, biểu đồ"],
        ["Người dùng", "Tìm kiếm, lọc, cấm/mở khóa tài khoản"],
        ["Duyệt KOL", "Duyệt hoặc từ chối hồ sơ KOL mới"],
        ["Duyệt Brand", "Duyệt hoặc từ chối hồ sơ thương hiệu mới"],
        ["Danh mục", "Quản lý danh mục ngành (tạo, sửa, xóa cây danh mục)"],
        ["Hoa hồng", "Xem tỷ lệ hoa hồng, ví nền tảng, doanh thu phí"],
        ["Rút tiền", "Duyệt yêu cầu rút tiền của KOL"],
    ])

    add_heading(doc, "5.2. Tổng quan nền tảng", 2)
    add_steps(doc, [
        "Vào \"Tổng quan\" (/admin).",
        "Xem các chỉ số: Tổng người dùng, KOL, Brand, Tổng booking, Doanh thu, Booking đang chạy.",
        "Chọn khoảng thời gian: 30 ngày / 90 ngày / 1 năm / Tùy chọn.",
        "Xem biểu đồ doanh thu và danh sách KOL hiệu quả nhất.",
    ])

    add_heading(doc, "5.3. Quản lý người dùng", 2)
    add_steps(doc, [
        "Vào \"Người dùng\".",
        "Tìm theo email; lọc theo vai trò (Admin / Brand / KOL) và trạng thái (Hoạt động / Bị cấm / Chờ xác minh).",
        "Với tài khoản vi phạm: bấm \"Cấm\" để khóa.",
        "Để mở lại: bấm \"Mở khóa\".",
    ])

    add_heading(doc, "5.4. Duyệt hồ sơ KOL", 2)
    add_steps(doc, [
        "Vào \"Duyệt KOL\".",
        "Xem danh sách hồ sơ \"Chờ duyệt\".",
        "Mở từng hồ sơ: kiểm tra thông tin, kênh MXH, gói giá, portfolio.",
        "Phù hợp → bấm \"Duyệt\" (KOL xuất hiện trên trang Khám phá).",
        "Không phù hợp → bấm \"Từ chối\", nhập lý do (tối thiểu 5 ký tự).",
    ])

    add_heading(doc, "5.5. Duyệt hồ sơ Brand", 2)
    add_steps(doc, [
        "Vào \"Duyệt Brand\".",
        "Xem hồ sơ công ty chờ duyệt — có thể mở chi tiết trước khi quyết định.",
        "Duyệt hoặc Từ chối kèm lý do (tương tự KOL).",
        "Brand chỉ đặt KOL được sau khi được duyệt.",
    ])

    add_heading(doc, "5.6. Quản lý danh mục", 2)
    add_steps(doc, [
        "Vào \"Danh mục\".",
        "Xem cây danh mục (mở/đóng từng nhánh).",
        "Thêm danh mục gốc hoặc danh mục con.",
        "Sửa tên, slug; xóa danh mục không còn dùng.",
    ])

    add_heading(doc, "5.7. Xử lý rút tiền KOL", 2)
    add_steps(doc, [
        "Vào \"Rút tiền\".",
        "Tab \"Chờ duyệt\": xem yêu cầu mới.",
        "Kiểm tra số tiền và thông tin ngân hàng.",
        "Bấm \"Duyệt\" nếu hợp lệ.",
        "Chuyển khoản thủ công cho KOL qua ngân hàng.",
        "Quay lại website → bấm \"Đã chi trả\" để hoàn tất.",
        "Nếu từ chối: bấm \"Từ chối\" và ghi lý do.",
    ])
    add_note(doc, "Luồng trạng thái rút tiền: Chờ duyệt → Đã duyệt → Đã chi trả (hoặc Từ chối).")
    doc.add_page_break()

    # 6. Chung
    add_heading(doc, "6. Các tính năng dùng chung", 1)

    add_heading(doc, "6.1. Chi tiết đơn đặt và trò chuyện", 2)
    doc.add_paragraph(
        "Khi mở một đơn đặt, bạn thấy 2 tab:"
    )
    add_bullets(doc, [
        "Chi tiết — thông tin chiến dịch, timeline trạng thái, các nút hành động",
        "Trò chuyện — nhắn tin trực tiếp với Brand (nếu bạn là KOL) hoặc KOL (nếu bạn là Brand)",
    ])

    add_heading(doc, "6.2. Thông báo", 2)
    doc.add_paragraph(
        "Biểu tượng chuông ở góc trên (Brand và KOL) hiển thị thông báo mới: "
        "đơn mới, thanh toán thành công, hồ sơ được duyệt, ứng tuyển mới... "
        "Vào trang \"Thông báo\" để xem tất cả, lọc chưa đọc."
    )

    add_heading(doc, "6.3. Đánh giá và nhận xét", 2)
    doc.add_paragraph(
        "Sau khi đơn \"Hoàn thành\", cả Brand và KOL có thể vào \"Đánh giá & Nhận xét\" "
        "để xem và để lại đánh giá (sao + nhận xét) cho đối tác."
    )

    # 7. FAQ
    add_heading(doc, "7. Câu hỏi thường gặp (FAQ)", 1)

    faqs = [
        ("Tôi quên mật khẩu, làm sao?", "Hiện tại liên hệ bộ phận hỗ trợ hoặc Admin để được reset mật khẩu."),
        ("Tại sao tôi không đặt được KOL?", "Kiểm tra: (1) Đã đăng nhập tài khoản Thương hiệu chưa? (2) Hồ sơ công ty đã được Admin duyệt chưa?"),
        ("Tại sao tôi không ứng tuyển được chiến dịch?", "Hồ sơ KOL phải ở trạng thái \"Đã duyệt\". Vào Trang quản lý KOL kiểm tra trạng thái."),
        ("Thanh toán VNPay báo lỗi?", "Thử lại sau vài phút. Kiểm tra số dư thẻ/ví. Nếu vẫn lỗi, chụp màn hình và liên hệ hỗ trợ."),
        ("KOL nộp deliverable rồi nhưng tôi chưa thấy?", "Vào chi tiết đơn, tab Chi tiết — kiểm tra trạng thái \"Đã giao\". Refresh trang nếu cần."),
        ("Tiền rút bao lâu về tài khoản?", "Sau khi Admin duyệt và đánh dấu \"Đã chi trả\", tiền chuyển theo thời gian xử lý ngân hàng (thường 1–3 ngày làm việc)."),
        ("Phiên đăng nhập hết hạn?", "Đăng nhập lại. Tránh để trang mở quá lâu không thao tác."),
        ("Tôi bị từ chối hồ sơ, làm gì tiếp?", "Đọc lý do từ chối, chỉnh sửa hồ sơ theo góp ý, rồi gửi duyệt lại."),
    ]
    for q, a in faqs:
        p = doc.add_paragraph()
        r = p.add_run(f"H: {q}")
        r.bold = True
        doc.add_paragraph(f"Đ: {a}")

    # 8. Glossary
    add_heading(doc, "8. Bảng thuật ngữ", 1)
    add_table(doc, ["Thuật ngữ", "Giải thích đơn giản"], [
        ["KOL", "Key Opinion Leader — người có ảnh hưởng trên mạng xã hội, làm nội dung quảng bá"],
        ["Brand / Thương hiệu", "Công ty, doanh nghiệp muốn thuê KOL quảng bá"],
        ["Chiến dịch", "Dự án marketing cần KOL tham gia (đăng bài, quay video...)"],
        ["Đơn đặt / Booking", "Thỏa thuận hợp tác cụ thể giữa Brand và KOL"],
        ["Deliverable", "Sản phẩm bàn giao — link bài đăng, video KOL đã đăng"],
        ["Escrow / Tạm giữ", "Tiền Brand trả được giữ an toàn cho đến khi xác nhận hoàn thành"],
        ["VNPay", "Cổng thanh toán trực tuyến tại Việt Nam"],
        ["Portfolio", "Tập hợp các bài đăng tiêu biểu của KOL"],
        ["Shortlist", "Danh sách ứng viên được Brand đánh dấu ưu tiên"],
        ["Hoa hồng", "Phí nền tảng thu trên mỗi giao dịch thành công"],
    ])

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("— Hết tài liệu —\nKOL Hub | Hướng dẫn sử dụng v1.0")
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    build_doc()
